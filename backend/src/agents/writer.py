import json
import os
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional

from jinja2 import Environment, BaseLoader
from langchain_ollama import ChatOllama
from langchain_core.prompts import ChatPromptTemplate


# ============================================================================
# Plantilla Jinja2 del documento de requisitos
# ============================================================================

PLANTILLA_MD = """# {{ nombre_proyecto }} — Especificación de Requisitos de Software

**Versión:** 1.0 · **Fecha:** {{ metadata.fecha }} · **Generado con MoSCoW AI** · **Requisitos:** {{ todos_requisitos | length }}

---

## Tabla de Contenidos

| # | Sección |
|---|---------|
| 1 | [Introducción](#1-introducción) |
| 2 | [Descripción General del Sistema](#2-descripción-general-del-sistema) |
| 3 | [Interesados del Proyecto](#3-interesados-del-proyecto) |
| 4 | [Requisitos Funcionales (RF)](#4-requisitos-funcionales-rf) |
| 5 | [Requisitos No Funcionales (RNF)](#5-requisitos-no-funcionales-rnf) |
| 6 | [Restricciones de Dominio (RD)](#6-restricciones-de-dominio-rd) |
| 7 | [Clasificación MoSCoW](#7-clasificación-moscow) |
| 8 | [Anexo — Razonamiento de Clasificación](#8-anexo--razonamiento-de-clasificación) |

---

## 1. Introducción

### 1.1 Propósito del Documento

{{ introduccion.proposito }}

### 1.2 Alcance del Sistema

{{ introduccion.alcance }}

### 1.3 Definiciones y Acrónimos

| Acrónimo | Definición |
|----------|-----------|
| RF | Requisito Funcional |
| RNF | Requisito No Funcional |
| RD | Restricción de Dominio |
| MoSCoW | Must Have / Should Have / Could Have / Won't Have |
| ISO 29148 | Estándar internacional para ingeniería de requisitos |

---

## 2. Descripción General del Sistema

{{ descripcion_general }}

---

## 3. Interesados del Proyecto
{% if stakeholders %}
| Nombre | Rol | Responsabilidades |
|--------|-----|-------------------|
{% for s in stakeholders %}| **{{ s.nombre }}** | {{ s.rol }} | {{ s.responsabilidades }} |
{% endfor %}
{% else %}
*No se registraron interesados.*
{% endif %}

---

## 4. Requisitos Funcionales (RF)
{% if requisitos_funcionales %}
{% for req in requisitos_funcionales %}
### {{ req.id }}

| Campo | Detalle |
|-------|---------|
| **Descripción** | {{ req.descripcion }} |
| **Prioridad** | {{ req.prioridad }} |
| **Criterio de Aceptación** | {{ req.criterio_aceptacion }} |
| **Clasificación MoSCoW** | {{ moscow_labels.get(req.id, '—') }} |

{% endfor %}
{% else %}
*No se registraron requisitos funcionales.*
{% endif %}

---

## 5. Requisitos No Funcionales (RNF)
{% if requisitos_no_funcionales %}
{% for req in requisitos_no_funcionales %}
### {{ req.id }}

| Campo | Detalle |
|-------|---------|
| **Descripción** | {{ req.descripcion }} |
| **Prioridad** | {{ req.prioridad }} |
| **Criterio de Aceptación** | {{ req.criterio_aceptacion }} |
| **Clasificación MoSCoW** | {{ moscow_labels.get(req.id, '—') }} |

{% endfor %}
{% else %}
*No se registraron requisitos no funcionales.*
{% endif %}

---

## 6. Restricciones de Dominio (RD)
{% if requisitos_dominio %}
{% for req in requisitos_dominio %}
### {{ req.id }}

| Campo | Detalle |
|-------|---------|
| **Descripción** | {{ req.descripcion }} |
| **Prioridad** | {{ req.prioridad }} |
| **Criterio de Aceptación** | {{ req.criterio_aceptacion }} |
| **Clasificación MoSCoW** | {{ moscow_labels.get(req.id, '—') }} |

{% endfor %}
{% else %}
*No se registraron restricciones de dominio.*
{% endif %}

---

## 7. Clasificación MoSCoW
{% if moscow_table %}
| ID | Descripción | Categoría | Score | Justificación |
|----|-------------|-----------|-------|---------------|
{% for row in moscow_table %}| {{ row.id }} | {{ row.descripcion[:55] }}{% if row.descripcion|length > 55 %}...{% endif %} | **{{ row.categoria }}** | {{ row.score }} | {{ row.justificacion[:90] }}{% if row.justificacion|length > 90 %}...{% endif %} |
{% endfor %}
{% else %}
| ID | Descripción | Prioridad |
|----|-------------|-----------|
{% for req in todos_requisitos %}| {{ req.id }} | {{ req.descripcion[:60] }}{% if req.descripcion|length > 60 %}...{% endif %} | {{ req.prioridad }} |
{% endfor %}
{% endif %}

---

## Anexo — Razonamiento de Clasificación

{% for req in todos_requisitos %}
### {{ req.id }} — {{ req.tipo }}

**Descripción:** {{ req.descripcion }}

**Prioridad declarada:** {{ req.prioridad }}
**Criterio de Aceptación:** {{ req.criterio_aceptacion }}

---
{% endfor %}

*Generado con MoSCoW AI · {{ metadata.fecha }}*
"""

CSS_PDF = """
@page { margin: 2.2cm 2.8cm; }
body {
    font-family: 'Segoe UI', Arial, sans-serif;
    font-size: 10.5pt;
    line-height: 1.65;
    color: #1e1e2e;
    background: #ffffff;
}
h1 { font-size: 20pt; color: #4f46e5; font-weight: 800; border-bottom: 2px solid #4f46e5; padding-bottom: 6px; margin-bottom: 4px; }
h2 { font-size: 13pt; color: #3730a3; font-weight: 700; border-bottom: 1px solid #e0e7ff; padding-bottom: 3px; margin-top: 24px; }
h3 { font-size: 10.5pt; color: #4338ca; font-weight: 600; margin-top: 16px; margin-bottom: 4px; }
p { margin: 6px 0; }
table { width: 100%; border-collapse: collapse; margin: 12px 0; font-size: 9pt; }
th { background-color: #4f46e5; color: #ffffff; padding: 7px 10px; text-align: left; font-weight: 600; }
td { padding: 6px 10px; border: 1px solid #e0e7ff; vertical-align: top; }
tr:nth-child(even) td { background-color: #f5f3ff; }
code { background: #ede9fe; color: #4f46e5; padding: 1px 5px; border-radius: 3px; font-family: 'Consolas', monospace; font-size: 9pt; }
hr { border: none; border-top: 1px solid #e0e7ff; margin: 20px 0; }
em { color: #7c3aed; font-style: italic; }
strong { color: #1e1e2e; font-weight: 700; }
li { margin: 3px 0; }
"""


# ============================================================================
# Agente Writer
# ============================================================================

class AgenteWriter:
    """
    Genera el documento formal de requisitos en MD, PDF y DOCX.
    Usa Jinja2 para renderizado, WeasyPrint para PDF y python-docx para DOCX.
    """

    def __init__(self):
        self.llm = ChatOllama(model="llama3.2", temperature=0.3)
        self.jinja_env = Environment(loader=BaseLoader())

    def _generar_introduccion(self, requisitos: List[Dict], nombre_proyecto: str = "", stakeholders_raw: str = "") -> Dict[str, str]:
        """Usa el LLM para generar la introducción narrativa del documento."""
        tipos = set(r.get("tipo", "") for r in requisitos)
        ids = [r.get("id", "") for r in requisitos]
        descripciones = [r.get("descripcion", "")[:80] for r in requisitos[:5]]

        prompt = ChatPromptTemplate.from_template(
            """Eres un Arquitecto de Software senior redactando un documento formal de especificación de requisitos.

            Proyecto: {nombre_proyecto}
            Requisitos ({total} en total):
            - IDs: {ids}
            - Tipos: {tipos}
            - Ejemplos de descripciones: {descripciones}
            {stakeholders_context}

            Genera exactamente un JSON con estas tres claves:
            {{
                "proposito": "Párrafo de 2-3 oraciones formales describiendo el propósito del documento y su audiencia objetivo",
                "alcance": "Párrafo de 2-3 oraciones describiendo qué cubre y qué excluye el sistema",
                "descripcion_general": "Párrafo de 3-4 oraciones describiendo el sistema, su contexto de negocio y valor aportado"
            }}

            Responde ÚNICAMENTE con el JSON, sin texto adicional.
            """
        )

        stakeholders_context = f"\nInteresados identificados: {stakeholders_raw}" if stakeholders_raw else ""

        chain = prompt | self.llm
        respuesta = chain.invoke({
            "nombre_proyecto": nombre_proyecto or "Sistema de Software",
            "ids": ", ".join(ids),
            "tipos": ", ".join(tipos),
            "total": len(requisitos),
            "descripciones": "; ".join(descripciones),
            "stakeholders_context": stakeholders_context,
        })

        try:
            raw = respuesta.content
            if isinstance(raw, str):
                contenido = raw.strip()
            else:
                item = raw[0]
                contenido = (item if isinstance(item, str) else item.get("text", "")).strip()
            if "```json" in contenido:
                contenido = contenido.split("```json")[1].split("```")[0].strip()
            elif "```" in contenido:
                contenido = contenido.split("```")[1].split("```")[0].strip()
            return json.loads(contenido)
        except Exception:
            return {
                "proposito": f"Este documento especifica los requisitos del sistema {nombre_proyecto}.",
                "alcance": "El sistema abarca las funcionalidades descritas en los requisitos listados a continuación.",
                "descripcion_general": "El sistema ha sido diseñado para satisfacer las necesidades identificadas durante el proceso de ingeniería de requisitos."
            }

    def _parsear_stakeholders(self, stakeholders_raw: str) -> List[Dict]:
        """Convierte texto libre de interesados en lista estructurada."""
        if not stakeholders_raw or not stakeholders_raw.strip():
            return []
        result = []
        for line in stakeholders_raw.strip().splitlines():
            line = line.strip()
            if not line:
                continue
            parts = [p.strip() for p in line.split("—") if p.strip()]
            if not parts:
                parts = [p.strip() for p in line.split("-", 2) if p.strip()]
            if len(parts) >= 3:
                result.append({"nombre": parts[0], "rol": parts[1], "responsabilidades": parts[2]})
            elif len(parts) == 2:
                result.append({"nombre": parts[0], "rol": parts[1], "responsabilidades": "—"})
            else:
                result.append({"nombre": parts[0], "rol": "—", "responsabilidades": "—"})
        return result

    def _clasificar_requisitos(self, requisitos: List[Dict]) -> Dict[str, List]:
        """Clasifica los requisitos por tipo."""
        funcionales, no_funcionales, dominio = [], [], []
        for req in requisitos:
            tipo = req.get("tipo", "").lower()
            if "no funcional" in tipo:
                no_funcionales.append(req)
            elif "dominio" in tipo:
                dominio.append(req)
            else:
                funcionales.append(req)
        return {"funcionales": funcionales, "no_funcionales": no_funcionales, "dominio": dominio}

    def _construir_moscow_table(self, requisitos: List[Dict], priorizacion: Optional[Dict]) -> List[Dict]:
        """Construye la tabla MoSCoW si hay datos de priorización disponibles."""
        if not priorizacion:
            return []

        priorizados = {
            r["requisito_id"]: r
            for r in priorizacion.get("requisitos_priorizados", [])
        }

        tabla = []
        for req in requisitos:
            req_id = req.get("id", "")
            p = priorizados.get(req_id, {})
            tabla.append({
                "id": req_id,
                "descripcion": req.get("descripcion", ""),
                "categoria": p.get("categoria_moscow", req.get("prioridad", "—")),
                "score": str(p.get("score_final", "—")),
                "justificacion": p.get("justificacion", "Sin datos de priorización")
            })

        return sorted(tabla, key=lambda x: x["score"], reverse=True)

    def _renderizar_markdown(self, requisitos: List[Dict], priorizacion: Optional[Dict] = None,
                              nombre_proyecto: str = "", stakeholders_raw: str = "") -> str:
        """Renderiza el documento completo en Markdown usando Jinja2."""
        clasificados = self._clasificar_requisitos(requisitos)
        textos = self._generar_introduccion(requisitos, nombre_proyecto, stakeholders_raw)
        moscow_table = self._construir_moscow_table(requisitos, priorizacion)
        stakeholders = self._parsear_stakeholders(stakeholders_raw)

        # Build moscow_labels dict for inline display in requirement tables
        moscow_labels: Dict[str, str] = {}
        if priorizacion:
            for r in priorizacion.get("requisitos_priorizados", []):
                moscow_labels[r["requisito_id"]] = r.get("categoria_moscow", "—")

        template = self.jinja_env.from_string(PLANTILLA_MD)
        return template.render(
            metadata={"fecha": datetime.now().strftime("%Y-%m-%d %H:%M:%S")},
            nombre_proyecto=nombre_proyecto or "Sistema de Software",
            introduccion={
                "proposito": textos.get("proposito", ""),
                "alcance": textos.get("alcance", "")
            },
            descripcion_general=textos.get("descripcion_general", ""),
            stakeholders=stakeholders,
            requisitos_funcionales=clasificados["funcionales"],
            requisitos_no_funcionales=clasificados["no_funcionales"],
            requisitos_dominio=clasificados["dominio"],
            todos_requisitos=requisitos,
            moscow_table=moscow_table,
            moscow_labels=moscow_labels,
        )

    def _exportar_pdf(self, contenido_md: str, ruta_pdf: str) -> bool:
        """Convierte el Markdown a PDF usando markdown2 + xhtml2pdf (o WeasyPrint como fallback)."""
        try:
            import markdown2
            html_body = markdown2.markdown(
                contenido_md,
                extras=["tables", "fenced-code-blocks", "header-ids"]
            )
            html_completo = (
                "<!DOCTYPE html><html lang='es'>"
                "<head><meta charset='utf-8'><title>Documento de Requisitos</title>"
                f"<style>{CSS_PDF}</style></head>"
                f"<body>{html_body}</body></html>"
            )
        except ImportError:
            print("[Writer] markdown2 no disponible — omitiendo PDF")
            return False

        # Intentar xhtml2pdf primero (Python puro, sin dependencias nativas)
        try:
            from xhtml2pdf import pisa
            with open(ruta_pdf, "wb") as f:
                resultado = pisa.CreatePDF(html_completo, dest=f, encoding="utf-8")  # type: ignore[assignment]
            if getattr(resultado, "err", None):
                raise RuntimeError(f"xhtml2pdf error code {resultado.err}")  # type: ignore[union-attr]
            return True
        except ImportError:
            pass
        except Exception as e:
            print(f"[Writer] xhtml2pdf falló: {e}")
            return False

        # Fallback: WeasyPrint (requiere GTK en Windows)
        try:
            from weasyprint import HTML, CSS
            HTML(string=html_completo).write_pdf(ruta_pdf, stylesheets=[CSS(string=CSS_PDF)])
            return True
        except ImportError:
            print("[Writer] xhtml2pdf y WeasyPrint no disponibles — omitiendo PDF")
        except Exception as e:
            print(f"[Writer] WeasyPrint falló: {e}")
        return False

    def _exportar_docx(self, requisitos: List[Dict], priorizacion: Optional[Dict], ruta_docx: str,
                       nombre_proyecto: str = "", stakeholders_raw: str = "") -> bool:
        """Genera el documento DOCX con python-docx."""
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            titulo = doc.add_heading(f"{nombre_proyecto or 'Sistema de Software'} — Especificación de Requisitos", 0)
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph("Generado con MoSCoW AI")
            doc.add_page_break()

            # Sección 1 — Introducción
            textos = self._generar_introduccion(requisitos, nombre_proyecto, stakeholders_raw)
            doc.add_heading("1. Introducción", 1)
            doc.add_heading("1.1 Propósito del Documento", 2)
            doc.add_paragraph(textos.get("proposito", ""))
            doc.add_heading("1.2 Alcance del Sistema", 2)
            doc.add_paragraph(textos.get("alcance", ""))
            doc.add_heading("1.3 Definiciones y Acrónimos", 2)
            t = doc.add_table(rows=1, cols=2)
            t.style = "Table Grid"
            t.rows[0].cells[0].text = "Acrónimo"
            t.rows[0].cells[1].text = "Definición"
            for acr, defn in [
                ("RF", "Requisito Funcional"),
                ("RNF", "Requisito No Funcional"),
                ("RD", "Requisito de Dominio"),
                ("MoSCoW", "Must Have / Should Have / Could Have / Won't Have"),
                ("ISO 29148", "Estándar internacional para ingeniería de requisitos"),
            ]:
                f = t.add_row().cells
                f[0].text = acr
                f[1].text = defn

            # Sección 2 — Descripción general
            doc.add_heading("2. Descripción General", 1)
            doc.add_paragraph(textos.get("descripcion_general", ""))

            # Sección 3 — Interesados
            doc.add_heading("3. Interesados del Proyecto", 1)
            stakeholders = self._parsear_stakeholders(stakeholders_raw)
            if stakeholders:
                tbl_s = doc.add_table(rows=1, cols=3)
                tbl_s.style = "Table Grid"
                enc = tbl_s.rows[0].cells
                enc[0].text = "Nombre"
                enc[1].text = "Rol"
                enc[2].text = "Responsabilidades"
                for s in stakeholders:
                    f = tbl_s.add_row().cells
                    f[0].text = s["nombre"]
                    f[1].text = s["rol"]
                    f[2].text = s["responsabilidades"]
            else:
                doc.add_paragraph("No se registraron interesados.")

            # Secciones 4-6 por tipo
            clasificados = self._clasificar_requisitos(requisitos)
            secciones = [
                ("4. Requisitos Funcionales (RF)", clasificados["funcionales"]),
                ("5. Requisitos No Funcionales (RNF)", clasificados["no_funcionales"]),
                ("6. Restricciones de Dominio (RD)", clasificados["dominio"]),
            ]
            for titulo_sec, lista in secciones:
                doc.add_heading(titulo_sec, 1)
                if not lista:
                    doc.add_paragraph("No se registraron requisitos en esta categoría.")
                    continue
                tbl = doc.add_table(rows=1, cols=4)
                tbl.style = "Table Grid"
                enc = tbl.rows[0].cells
                enc[0].text = "ID"
                enc[1].text = "Descripción"
                enc[2].text = "Prioridad"
                enc[3].text = "Criterio de Aceptación"
                for req in lista:
                    fila = tbl.add_row().cells
                    fila[0].text = req.get("id", "")
                    fila[1].text = req.get("descripcion", "")
                    fila[2].text = req.get("prioridad", "")
                    fila[3].text = req.get("criterio_aceptacion", "")

            # Sección 7 — MoSCoW
            doc.add_heading("7. Clasificación MoSCoW", 1)
            moscow_table = self._construir_moscow_table(requisitos, priorizacion)
            if moscow_table:
                tbl_m = doc.add_table(rows=1, cols=5)
                tbl_m.style = "Table Grid"
                enc = tbl_m.rows[0].cells
                enc[0].text = "ID"
                enc[1].text = "Descripción"
                enc[2].text = "Categoría"
                enc[3].text = "Score"
                enc[4].text = "Justificación"
                for row in moscow_table:
                    f = tbl_m.add_row().cells
                    f[0].text = row["id"]
                    f[1].text = row["descripcion"][:80]
                    f[2].text = row["categoria"]
                    f[3].text = str(row["score"])
                    f[4].text = row["justificacion"][:120]
            else:
                tbl_m = doc.add_table(rows=1, cols=3)
                tbl_m.style = "Table Grid"
                enc = tbl_m.rows[0].cells
                enc[0].text = "ID"
                enc[1].text = "Descripción"
                enc[2].text = "Prioridad"
                for req in requisitos:
                    f = tbl_m.add_row().cells
                    f[0].text = req.get("id", "")
                    f[1].text = req.get("descripcion", "")[:80]
                    f[2].text = req.get("prioridad", "")

            # Anexo B
            doc.add_page_break()
            doc.add_heading("ANEXO B — Razonamiento de Clasificación", 1)
            for i, req in enumerate(requisitos, 1):
                doc.add_heading(f"B.{i} — {req.get('id', '')}", 2)
                doc.add_paragraph(f"Descripción: {req.get('descripcion', '')}")
                doc.add_paragraph(f"Tipo: {req.get('tipo', '')}")
                doc.add_paragraph(f"Prioridad: {req.get('prioridad', '')}")
                doc.add_paragraph(f"Criterio de Aceptación: {req.get('criterio_aceptacion', '')}")

            doc.save(ruta_docx)
            return True
        except ImportError:
            print("[Writer] python-docx no disponible — omitiendo DOCX")
            return False
        except Exception as e:
            print(f"[Writer] Error generando DOCX: {e}")
            return False

    def generar_documento(self, archivo_requisitos: str, archivo_priorizacion: Optional[str] = None,
                          output_dir: Optional[str] = None, nombre_proyecto: str = "",
                          stakeholders_raw: str = "") -> Dict[str, Any]:
        """
        Genera el documento formal en MD, PDF y DOCX.

        Args:
            archivo_requisitos: Ruta al JSON con la lista de requisitos.
            archivo_priorizacion: Ruta opcional al JSON de priorización MoSCoW.

        Returns:
            Dict con rutas a los archivos generados (md, pdf, docx).
        """
        with open(archivo_requisitos, "r", encoding="utf-8") as f:
            data = json.load(f)

        if isinstance(data, list):
            requisitos = data
        elif isinstance(data, dict) and "requisitos" in data:
            requisitos = data["requisitos"]
        else:
            requisitos = [data]

        if not requisitos:
            raise ValueError("No hay requisitos para generar el documento")

        # Intentar cargar priorización automáticamente si existe
        priorizacion = None
        if archivo_priorizacion and os.path.exists(archivo_priorizacion):
            with open(archivo_priorizacion, "r", encoding="utf-8") as f:
                priorizacion = json.load(f)
        elif os.path.exists("outputs/requisitos_priorizados.json"):
            with open("outputs/requisitos_priorizados.json", "r", encoding="utf-8") as f:
                priorizacion = json.load(f)

        directorio = Path(output_dir) if output_dir else Path("outputs/documento_requisitos")
        directorio.mkdir(parents=True, exist_ok=True)

        print("\n[Writer] Generando documento formal de requisitos...")

        # Paso 1: Markdown con Jinja2
        print("[Writer] Renderizando Markdown con Jinja2...")
        contenido_md = self._renderizar_markdown(requisitos, priorizacion, nombre_proyecto, stakeholders_raw)
        ruta_md = directorio / "documento_requisitos.md"
        ruta_md.write_text(contenido_md, encoding="utf-8")
        print(f"[Writer] MD generado: {ruta_md}")

        # Paso 2: PDF con WeasyPrint
        ruta_pdf_str = str(directorio / "documento_requisitos.pdf")
        print("[Writer] Convirtiendo a PDF con WeasyPrint...")
        ruta_pdf = ruta_pdf_str if self._exportar_pdf(contenido_md, ruta_pdf_str) else None

        # Paso 3: DOCX con python-docx
        ruta_docx_str = str(directorio / "documento_requisitos.docx")
        print("[Writer] Generando DOCX con python-docx...")
        ruta_docx = ruta_docx_str if self._exportar_docx(requisitos, priorizacion, ruta_docx_str, nombre_proyecto, stakeholders_raw) else None

        resultado = {
            "md": str(ruta_md),
            "pdf": ruta_pdf,
            "docx": ruta_docx,
            "total_requisitos": len(requisitos),
            "fecha": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }

        print(f"[Writer] PDF:  {resultado['pdf'] or 'no disponible'}")
        print(f"[Writer] DOCX: {resultado['docx'] or 'no disponible'}")

        return resultado
